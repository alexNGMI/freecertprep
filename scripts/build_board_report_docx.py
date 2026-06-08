from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "board-report-2026-06-08.md"
OUTPUT = ROOT / "docs" / "freecertprep-board-report-2026-06-08.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 33, 45)
MUTED = RGBColor(92, 103, 116)
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"
RISK = RGBColor(155, 28, 28)
GREEN = RGBColor(31, 105, 70)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_run_font(run, size=None, color=INK, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline(paragraph, text, size=11, color=INK):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size - 0.5, color=DARK_BLUE, font="Consolas")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, color=color)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_section(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("freecertprep | Board Review | June 8, 2026")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    add_page_field(footer.paragraphs[0])


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def add_table(doc, rows):
    if not rows:
        return
    columns = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"

    if columns == 2:
        widths = [2700, 6660]
    elif columns == 3:
        widths = [2500, 1800, 5060]
    elif columns == 4:
        widths = [2700, 1800, 1700, 3160]
    else:
        base = 9360 // columns
        widths = [base] * columns
        widths[-1] += 9360 - sum(widths)

    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value, size=9.5 if columns >= 4 else 10)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = DARK_BLUE
            elif row_index % 2 == 0:
                set_cell_shading(cell, "FAFBFC")

    mark_header_row(table.rows[0])
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("BOARD REVIEW")
    set_run_font(run, size=23, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("freecertprep Product, Content, Technology, and Strategy")
    set_run_font(run, size=15, color=MUTED)

    metadata = [
        ("To", "Board and strategic stakeholders"),
        ("Date", "June 8, 2026"),
        ("Scope", "Current offerings, recent changes, risk, roadmap, and board decisions"),
        ("Repository", "alexNGMI/freecertprep"),
        ("Status", "Decision and prioritization review"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=11, color=INK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=11, color=INK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(16)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Recommendation")
    set_run_font(run, size=12, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline(
        p,
        "Pause major catalog expansion. Clear security and trust debt, harden the weakest live banks, establish measurable user signals, and only then choose the next sister-site or backend investment.",
        size=12,
        color=INK,
    )
    doc.add_page_break()


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_cover(doc)

    index = 1
    paragraph_buffer = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        text = " ".join(line.strip() for line in paragraph_buffer).strip()
        if text:
            p = doc.add_paragraph()
            add_inline(p, text)
        paragraph_buffer = []

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{2,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1)) - 1
            title = heading.group(2)
            if level == 1 and title in {
                "Current Product Offering",
                "Recent Development Program",
                "Product and Technical Architecture",
                "Quality Review",
                "Key Risks",
                "Competitive and Strategic Position",
                "Recommended Roadmap",
            }:
                doc.add_page_break()
            doc.add_heading(title, level=level)
            index += 1
            continue

        if re.match(r"^-\s+", stripped):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^-\s+", "", stripped))
            index += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
            index += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            p_pr = p._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), PALE_BLUE)
            p_pr.append(shading)
            borders = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "18")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), "2E74B5")
            borders.append(left)
            p_pr.append(borders)
            add_inline(p, stripped[2:], size=11, color=DARK_BLUE)
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()

    core = doc.core_properties
    core.title = "freecertprep Board Review"
    core.subject = "Product, content, technology, risk, and strategy review"
    core.author = "freecertprep"
    core.keywords = "board review, certification prep, roadmap, product strategy"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
